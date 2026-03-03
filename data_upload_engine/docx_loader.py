"""
Word document (.docx) loader for Local Analyst.
Extracts text, tables, and structure from Word documents.
Integrates with existing loader pattern.
"""

import pandas as pd
from pathlib import Path
from typing import Tuple, Dict, Any, List, Optional
from dataclasses import dataclass


@dataclass
class DocxContent:
    """Structured content from Word document."""
    text_blocks: List[Dict[str, Any]]
    tables: List[pd.DataFrame]
    paragraphs: List[str]
    headings: List[Dict[str, str]]
    metadata: Dict[str, Any]


def load_docx(file_path: str) -> Tuple[DocxContent, Dict[str, Any]]:
    """
    Load Word document and extract all content.
    
    Args:
        file_path: Path to .docx file
        
    Returns:
        Tuple of (DocxContent, metadata dict)
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    
    path = Path(file_path)
    doc = Document(file_path)
    
    # Extract paragraphs
    paragraphs = []
    headings = []
    text_blocks = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
            
        # Detect headings
        style_name = para.style.name
        if 'Heading' in style_name:
            level = style_name.replace('Heading ', '')
            headings.append({
                'level': level,
                'text': text,
                'position': i
            })
        
        paragraphs.append(text)
        text_blocks.append({
            'paragraph_num': i + 1,
            'text': text,
            'style': style_name
        })
    
    # Extract tables
    tables = []
    for table_num, table in enumerate(doc.tables, 1):
        # Get all cells as 2D list
        data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            data.append(row_data)
        
        # Convert to DataFrame
        if data:
            # Try to use first row as headers if they look like headers
            first_row = data[0]
            if len(set(first_row)) == len(first_row) and all(first_row):
                # First row looks like headers
                df = pd.DataFrame(data[1:], columns=first_row)
            else:
                # Use default column names
                df = pd.DataFrame(data)
            
            tables.append(df)
    
    # Extract core properties (metadata)
    core_props = doc.core_properties
    metadata = {
        'source_file': path.name,
        'file_type': 'docx',
        'paragraphs': len(paragraphs),
        'tables': len(tables),
        'headings': len(headings),
        'author': core_props.author if hasattr(core_props, 'author') else None,
        'title': core_props.title if hasattr(core_props, 'title') else None,
        'subject': core_props.subject if hasattr(core_props, 'subject') else None,
        'created': core_props.created if hasattr(core_props, 'created') else None,
        'modified': core_props.modified if hasattr(core_props, 'modified') else None,
    }
    
    content = DocxContent(
        text_blocks=text_blocks,
        tables=tables,
        paragraphs=paragraphs,
        headings=headings,
        metadata=metadata
    )
    
    return content, metadata


def docx_to_dataframe(
    file_path: str,
    table_index: int = 0
) -> Tuple[pd.DataFrame, Dict[str, Any]]:
    """
    Extract specific table from Word document as DataFrame.
    
    Args:
        file_path: Path to .docx file
        table_index: Which table to extract (0-indexed)
        
    Returns:
        Tuple of (DataFrame, metadata dict)
    """
    content, metadata = load_docx(file_path)
    
    if not content.tables:
        # No tables - create DataFrame from text blocks
        if content.text_blocks:
            df = pd.DataFrame(content.text_blocks)
        else:
            df = pd.DataFrame()
        
        metadata['note'] = 'No tables found - returning text blocks'
        return df, metadata
    
    if table_index >= len(content.tables):
        # Requested table doesn't exist - return first table
        table_index = 0
        metadata['warning'] = f'Requested table {table_index} not found, using table 0'
    
    df = content.tables[table_index]
    metadata['table_index'] = table_index
    metadata['table_count'] = len(content.tables)
    
    return df, metadata


def get_docx_summary(file_path: str) -> Dict[str, Any]:
    """
    Get summary information about Word document without full parsing.
    
    Args:
        file_path: Path to .docx file
        
    Returns:
        Dict with summary info
    """
    content, metadata = load_docx(file_path)
    
    summary = {
        'file_name': Path(file_path).name,
        'paragraphs': metadata['paragraphs'],
        'tables': metadata['tables'],
        'headings': metadata['headings'],
        'total_text_length': sum(len(p) for p in content.paragraphs),
        'text_preview': '\n'.join(content.paragraphs[:3]),
        'heading_structure': content.headings,
        'table_shapes': [t.shape for t in content.tables],
        'has_tables': len(content.tables) > 0
    }
    
    # Add document properties if available
    if metadata.get('title'):
        summary['title'] = metadata['title']
    if metadata.get('author'):
        summary['author'] = metadata['author']
    if metadata.get('subject'):
        summary['subject'] = metadata['subject']
    
    return summary


def extract_text_only(file_path: str) -> str:
    """
    Extract all text from Word document as single string.
    
    Args:
        file_path: Path to .docx file
        
    Returns:
        Combined text content
    """
    content, _ = load_docx(file_path)
    return '\n\n'.join(content.paragraphs)


def extract_tables_as_list(file_path: str) -> List[pd.DataFrame]:
    """
    Extract all tables from Word document.
    
    Args:
        file_path: Path to .docx file
        
    Returns:
        List of DataFrames (one per table)
    """
    content, _ = load_docx(file_path)
    return content.tables


# Example usage
if __name__ == "__main__":
    # Test with a sample document
    sample_path = "sample.docx"
    
    # Get summary
    summary = get_docx_summary(sample_path)
    print(f"Document: {summary['file_name']}")
    print(f"Paragraphs: {summary['paragraphs']}")
    print(f"Tables: {summary['tables']}")
    print(f"Text preview:\n{summary['text_preview']}")
    
    # Extract first table
    df, metadata = docx_to_dataframe(sample_path, table_index=0)
    print(f"\nFirst table shape: {df.shape}")
    print(df.head())
